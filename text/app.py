from flask import Flask, render_template, request, send_file
import markdown
from pathlib import Path
import pdfplumber
from docx import Document
from bs4 import BeautifulSoup
import tempfile
import os
import re

app = Flask(__name__)

def clean_content(content):
    # 去除颜色格式（如\x1b[31m等）
    content = re.sub(r'\x1b\[[0-9;]*m', '', content)
    
    # 去除 Markdown 格式
    content = re.sub(r'<[^>]+>', '', content)  # 去除 HTML 标签
    content = re.sub(r'\[.*?\]\(.*?\)', '', content)  # 去除 Markdown 链接
    # 保留标题格式
    content = re.sub(r'#{1,6}\s+', '# ', content)  # 标准化标题格式
    content = re.sub(r'\*{1,2}(.*?)\*{1,2}', r'\1', content)  # 去除加粗/斜体
    content = re.sub(r'`{1,3}(.*?)`{1,3}', r'\1', content)  # 去除代码标记
    
    # 处理空格和空行
    content = re.sub(r' {2,}', ' ', content)  # 合并多个空格
    content = re.sub(r'\n{3,}', '\n\n', content)  # 多个空行合并为两个
    content = re.sub(r'\n\s*\n', '\n\n', content)  # 去除空行中的空格
    
    # 智能分段和标点处理
    paragraphs = content.split('\n')
    cleaned_paragraphs = []
    for para in paragraphs:
        if not para.strip():
            cleaned_paragraphs.append('')
            continue
            
        # 处理句子
        sentences = re.split(r'(?<=[。！？])', para)
        cleaned_sentences = []
        for sentence in sentences:
            sentence = sentence.strip()
            if not sentence:
                continue
                
            # 去除多余空格
            sentence = re.sub(r'\s+', ' ', sentence)
            
            # 自动添加标点
            if not re.search(r'[。！？]$', sentence):
                # 根据上下文判断标点类型
                if re.search(r'[吗呢吧]$', sentence):
                    sentence += '？'
                elif re.search(r'[了着过]$', sentence):
                    sentence += '。'
                else:
                    sentence += '。'
                    
            cleaned_sentences.append(sentence)
            
        cleaned_paragraphs.append(' '.join(cleaned_sentences))
    
    return '\n'.join(cleaned_paragraphs)

def process_file(file):
    filename = file.filename
    ext = os.path.splitext(filename)[1].lower()
    
    try:
        if ext == '.md':
            content = file.read().decode('utf-8')
        elif ext == '.docx':
            doc = Document(file)
            content = '\n'.join([p.text for p in doc.paragraphs])
        elif ext == '.pdf':
            with pdfplumber.open(file) as pdf:
                content = '\n'.join([page.extract_text() for page in pdf.pages])
        elif ext == '.html':
            soup = BeautifulSoup(file.read(), 'html.parser')
            content = soup.get_text()
        else:
            raise ValueError('不支持的文件格式')
        
        return clean_content(content)
    except Exception as e:
        print(f"文件处理错误: {str(e)}")
        raise

@app.route('/')
def index():
    return render_template('index.html')

@app.route('/upload', methods=['POST'])
def upload():
    if 'file' not in request.files:
        return 'No file uploaded', 400
    
    file = request.files['file']
    if file.filename == '':
        return 'No selected file', 400
    
    try:
        content = process_file(file)
        html_content = markdown.markdown(content)
        
        # Save to temp file
        doc = Document()
        doc.add_paragraph(content)
        
        with tempfile.NamedTemporaryFile(delete=False, suffix='.docx') as tmp:
            doc.save(tmp.name)
            tmp_path = tmp.name
        
        return {
            'preview': html_content,
            'download_url': f'/download/{os.path.basename(tmp_path)}'
        }
    except Exception as e:
        return str(e), 500

@app.route('/download/<filename>')
def download(filename):
    path = os.path.join(tempfile.gettempdir(), filename)
    return send_file(path, as_attachment=True)

if __name__ == '__main__':
    app.run(debug=True)
